# res = runRowQuery(145)
# print(len(res))

# res = runRowQuery(34)
# print(len(res))

# res = runRowQuery(576) # 5
# print(len(res))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from c_queryParser import rowsParsing, headerFooterParsing
from docx.shared import Pt


def createFooterFunction(docNumber):
    headerList, footerList = headerFooterParsing(docNumber)
    document = Document('vv.docx')
    style = document.styles['Normal']
    font = style.font
    font.name = 'Cascadia Code'
    font.complex_script = True
    font.rtl = True
    font.size = Pt(5)
    if len(document.tables) > 2:
        table = document.tables[2]
        for i in range(len(footerList)):
            for j in range(len(table.columns)):
                cell = table.cell(i, j)
                cell.text = str(footerList[i])
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if len(document.tables) > 0:
        header_table = document.tables[0]
    if len(header_table.rows) * len(header_table.columns) == len(headerList):
        for j, data in enumerate(headerList):
            row_index = j // len(header_table.columns)
            col_index = j % len(header_table.columns)
            header_cell = header_table.cell(row_index, col_index)
            header_cell.text = data
            for paragraph in header_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print("number of cells does not match the length of headerList.")
    document.save('lastPage.docx')


def createMainsPages(docNumber):
    finalDataList = rowsParsing(docNumber)
    document = Document('pt3.docx')
    style = document.styles['Normal']
    font = style.font
    font.name = 'rpt-Regular'
    font.complex_script = True
    font.rtl = True
    font.size = Pt(5)
    if len(document.tables) > 1:
        table = document.tables[1]
        num_rows = len(table.rows)
        if num_rows >= len(finalDataList):
            for j in range(len(finalDataList)):
                for k in range(len(finalDataList[j])):
                    cell = table.cell(j, k)
                    cell.margin_top = Pt(0)
                    cell.margin_right = Pt(0)
                    cell.text = finalDataList[j][k]
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save('vv.docx')
    createFooterFunction(576)


createMainsPages(576)
